import re
from pathlib import Path
from docx import Document
from typing import Set, List
from app.utils.validation import PlaceholderValidator
from app.core.exceptions import PlaceholderExtractionError
from app.utils.logging_config import app_logger

class PlaceholderExtractor:
    @staticmethod
    def extract_from_docx(file_path: Path) -> Set[str]:
        """Extract all {{placeholders}} from DOCX file"""
        try:
            doc = Document(file_path)
            placeholders = set()
            
            # Pattern to match {{placeholder}} with optional spaces
            pattern = r'\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}'
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                # Use direct regex instead of validator for better capture
                found = re.findall(pattern, text)
                placeholders.update(found)
                # Also try validator method as fallback
                found2 = PlaceholderValidator.extract_placeholders_from_text(text)
                placeholders.update(found2)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            found = re.findall(pattern, text)
                            placeholders.update(found)
                            found2 = PlaceholderValidator.extract_placeholders_from_text(text)
                            placeholders.update(found2)
            
            # Extract from headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        text = paragraph.text
                        found = re.findall(pattern, text)
                        placeholders.update(found)
                        found2 = PlaceholderValidator.extract_placeholders_from_text(text)
                        placeholders.update(found2)
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        text = paragraph.text
                        found = re.findall(pattern, text)
                        placeholders.update(found)
                        found2 = PlaceholderValidator.extract_placeholders_from_text(text)
                        placeholders.update(found2)
            
            app_logger.info(f"Extracted {len(placeholders)} placeholders: {placeholders}")
            return placeholders
            
        except Exception as e:
            app_logger.error(f"Failed to extract placeholders: {str(e)}")
            raise PlaceholderExtractionError(f"Cannot extract placeholders: {str(e)}")
    
    @staticmethod
    def validate_placeholders(placeholders: Set[str]) -> bool:
        """Validate placeholder naming conventions"""
        for placeholder in placeholders:
            # Check for valid characters
            if not re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*$', placeholder):
                raise PlaceholderExtractionError(
                    f"Invalid placeholder name: '{placeholder}'. "
                    f"Use only letters, numbers, underscores, starting with letter or underscore"
                )
        return True