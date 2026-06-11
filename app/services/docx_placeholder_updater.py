from docx import Document
from pathlib import Path
import re
from app.utils.logging_config import app_logger

def update_placeholder_in_docx(file_path: str, old_name: str, new_name: str) -> bool:
    """
    Update placeholder name in DOCX file.
    Replaces {{old_name}} with {{new_name}} in the document.
    """
    try:
        doc = Document(file_path)
        modified = False
        
        # Pattern to match {{old_name}} (with optional spaces)
        old_pattern = re.compile(r'\{\{\s*' + re.escape(old_name) + r'\s*\}\}')
        new_placeholder = f"{{{{{new_name}}}}}"
        
        # Update in paragraphs
        for paragraph in doc.paragraphs:
            if old_pattern.search(paragraph.text):
                paragraph.text = old_pattern.sub(new_placeholder, paragraph.text)
                modified = True
        
        # Update in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_pattern.search(paragraph.text):
                            paragraph.text = old_pattern.sub(new_placeholder, paragraph.text)
                            modified = True
        
        # Update in headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if old_pattern.search(paragraph.text):
                        paragraph.text = old_pattern.sub(new_placeholder, paragraph.text)
                        modified = True
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if old_pattern.search(paragraph.text):
                        paragraph.text = old_pattern.sub(new_placeholder, paragraph.text)
                        modified = True
        
        if modified:
            doc.save(file_path)
            app_logger.info(f"✅ Updated DOCX file: {file_path} (replaced '{old_name}' → '{new_name}')")
        else:
            app_logger.info(f"⚠️ No placeholder '{old_name}' found in DOCX file: {file_path}")
        
        return True
        
    except Exception as e:
        app_logger.error(f"Error updating DOCX {file_path}: {str(e)}")
        return False