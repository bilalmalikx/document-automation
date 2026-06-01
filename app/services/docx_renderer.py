from pathlib import Path
from docx import Document
from typing import Dict
import io
import tempfile
import shutil
from app.core.exceptions import RenderError
from app.utils.logging_config import app_logger

class DOCXRenderer:
    @staticmethod
    def render_template(
        template_path: Path,
        context: Dict[str, str]
    ) -> bytes:
        """
        Render DOCX - Preserve line length
        """
        temp_file = None
        try:
            # Create temp copy
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                shutil.copy2(template_path, tmp.name)
                temp_file = tmp.name
            
            # Load document
            doc = Document(temp_file)
            
            # Replace placeholders with padding to maintain length
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if not text:
                    continue
                
                new_text = text
                for key, value in context.items():
                    if value:
                        placeholder = "{{" + key + "}}"
                        if placeholder in new_text:
                            # Calculate padding to maintain length
                            placeholder_len = len(placeholder)
                            value_len = len(str(value))
                            
                            if value_len < placeholder_len:
                                # Add spaces to match original length
                                padding = " " * (placeholder_len - value_len)
                                replacement = str(value) + padding
                            else:
                                replacement = str(value)
                            
                            new_text = new_text.replace(placeholder, replacement)
                
                if new_text != text:
                    paragraph.text = new_text
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            if not text:
                                continue
                            
                            new_text = text
                            for key, value in context.items():
                                if value:
                                    placeholder = "{{" + key + "}}"
                                    if placeholder in new_text:
                                        placeholder_len = len(placeholder)
                                        value_len = len(str(value))
                                        
                                        if value_len < placeholder_len:
                                            padding = " " * (placeholder_len - value_len)
                                            replacement = str(value) + padding
                                        else:
                                            replacement = str(value)
                                        
                                        new_text = new_text.replace(placeholder, replacement)
                            
                            if new_text != text:
                                paragraph.text = new_text
            
            # Save
            output_stream = io.BytesIO()
            doc.save(output_stream)
            output_stream.seek(0)
            
            return output_stream.getvalue()
            
        except Exception as e:
            app_logger.error(f"Rendering failed: {str(e)}")
            raise RenderError(f"Cannot render document: {str(e)}")
        finally:
            if temp_file and Path(temp_file).exists():
                Path(temp_file).unlink()